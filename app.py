import os
import requests
import streamlit as st
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from PIL import Image
import pytesseract

load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"

st.set_page_config(page_title="AI Meeting Notes Summarizer", page_icon="🧠", layout="wide")

theme_choice = st.sidebar.radio("🎨 Theme", ["Light", "Dark"], index=0)

light_css = """
    <style>
    h1 { font-size: 42px !important; color: #2E86C1; }
    .stTextArea textarea { font-size: 16px !important; line-height: 1.5; }
    .stDownloadButton button {
        background-color: #2E86C1; color: white; border-radius: 8px; padding: 8px 16px; font-size: 16px;
    }
    .stDownloadButton button:hover { background-color: #1B4F72; }
    .stButton>button {
        background-color: #117A65; color: white; border-radius: 8px; padding: 10px 20px; font-size: 18px;
    }
    .stButton>button:hover { background-color: #0B5345; }
    .summary-box {
        font-size: 18px; background: #F4F6F7; padding: 15px; border-radius: 10px; line-height: 1.6;
        color: #111;
    }
    </style>
"""

dark_css = """
    <style>
    body, .stApp { background-color: #121212 !important; color: #EAECEE !important; }
    h1 { font-size: 42px !important; color: #5DADE2; }
    .stTextArea textarea { font-size: 16px !important; line-height: 1.5; background: #1C1C1C; color: #EAECEE; }
    .stDownloadButton button {
        background-color: #5DADE2; color: white; border-radius: 8px; padding: 8px 16px; font-size: 16px;
    }
    .stDownloadButton button:hover { background-color: #3498DB; }
    .stButton>button {
        background-color: #28B463; color: white; border-radius: 8px; padding: 10px 20px; font-size: 18px;
    }
    .stButton>button:hover { background-color: #1D8348; }
    .summary-box {
        font-size: 18px; background: #1E1E1E; padding: 15px; border-radius: 10px; line-height: 1.6;
        color: #EAECEE;
    }
    </style>
"""

st.markdown(dark_css if theme_choice == "Dark" else light_css, unsafe_allow_html=True)

def extract_text_from_file(file):
    """Extract text from .txt, .pdf, .png/.jpg images."""
    if file.name.endswith(".txt"):
        return file.read().decode("utf-8")

    elif file.name.endswith(".pdf"):
        reader = PdfReader(file)
        return "\n".join([page.extract_text() or "" for page in reader.pages])

    elif file.type.startswith("image/"):
        image = Image.open(file)
        return pytesseract.image_to_string(image)

    else:
        return ""

def call_groq_api(text: str, instruction: str) -> str:
    if not GROQ_API_KEY:
        raise RuntimeError("GROQ_API_KEY is not set.")
    headers = {"Authorization": f"Bearer {GROQ_API_KEY}", "Content-Type": "application/json"}
    payload = {
        "model": "llama-3.1-8b-instant",
        "messages": [
            {"role": "system", "content": "You are an AI summarizer that must strictly follow the instructions."},
            {"role": "user", "content": f"Instruction: {instruction}\n\nText:\n{text}"},
        ],
        "temperature": 0.5,
    }
    r = requests.post(GROQ_API_URL, headers=headers, json=payload, timeout=60)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]

def make_txt(s: str): return BytesIO(s.encode("utf-8"))
def make_md(s: str): return BytesIO(f"# Summary\n\n{s}".encode("utf-8"))
def make_docx(s: str):
    buf = BytesIO(); doc = Document()
    doc.add_heading("Meeting Summary", 0); doc.add_paragraph(s)
    doc.save(buf); buf.seek(0); return buf
def make_pdf(s: str):
    buf = BytesIO(); c = canvas.Canvas(buf, pagesize=letter)
    w, h = letter; t = c.beginText(50, h - 50); t.setFont("Helvetica", 12)
    for line in s.split("\n"): t.textLine(line)
    c.drawText(t); c.save(); buf.seek(0); return buf

if "history" not in st.session_state:
    st.session_state.history = []


st.markdown("<h1>AI Meeting Notes Summarizer & Sharer</h1>", unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    uploaded_files = st.file_uploader(
        "Upload Files (.txt, .pdf, .png, .jpg, .jpeg)", 
        type=["txt", "pdf", "png", "jpg", "jpeg"], 
        accept_multiple_files=True
    )
with col2:
    manual_text = st.text_area("Paste or write text here", height=200)

st.sidebar.header(" Summary Options")
instruction_options = st.sidebar.multiselect(
    "Choose one or more styles:",
    [
        "Summarize in bullet points",
        "Explain in simple terms",
        "Extract only key points",
        "Highlight only action items",
        "Detailed summary"
    ],
    default=["Summarize in bullet points"]
)
custom_instruction = st.sidebar.text_area("Or write your own instruction (optional)")

# Merge all instructions
final_instruction = "; ".join(instruction_options)
if custom_instruction.strip():
    final_instruction += f"; {custom_instruction}"

if st.button("🚀 Generate Summary"):
    text_data = ""

    if uploaded_files:
        for file in uploaded_files:
            extracted = extract_text_from_file(file)
            text_data += extracted + "\n"

    if manual_text.strip():
        text_data += "\n" + manual_text.strip()

    if not text_data.strip():
        st.warning("Please upload or enter some text.")
    else:
        try:
            with st.spinner("🤖 Generating summary..."):
                summary = call_groq_api(text_data, final_instruction or "Summarize the text in a structured way.")

            st.success("Summary Generated! You can edit it below:")

            st.markdown(f"<div class='summary-box'>{summary}</div>", unsafe_allow_html=True)

            editable = st.text_area("✍️ Edit Summary if needed", summary, height=400)

            # Save in history
            st.session_state.history.append({
                "instruction": final_instruction or "(default summary)",
                "summary": editable
            })

            st.subheader("⬇️ Download Summary")
            c1, c2, c3, c4 = st.columns(4)
            with c1: st.download_button("📄 TXT", make_txt(editable), "summary.txt", "text/plain")
            with c2: st.download_button("📝 Markdown", make_md(editable), "summary.md", "text/markdown")
            with c3: st.download_button("📘 DOCX", make_docx(editable), "summary.docx",
                                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with c4: st.download_button("📕 PDF", make_pdf(editable), "summary.pdf", "application/pdf")

        except Exception as e:
            st.error(f"Error: {e}")

# --- HISTORY ---
if st.session_state.history:
    st.sidebar.title("History")
    for i, h in enumerate(reversed(st.session_state.history), 1):
        with st.sidebar.expander(f"Summary {len(st.session_state.history) - i + 1}"):
            st.write(f"**Instruction(s):** {h['instruction']}")
            st.markdown(f"<div class='summary-box'>{h['summary']}</div>", unsafe_allow_html=True)
            st.download_button("📄 Download TXT", make_txt(h["summary"]), f"summary_{i}.txt", "text/plain",
                               key=f"hist_dl_{i}")
